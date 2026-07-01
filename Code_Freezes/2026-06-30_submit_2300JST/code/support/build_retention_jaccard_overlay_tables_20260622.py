#!/usr/bin/env python3
"""Build manuscript-facing Retention/Jaccard overlay quantification tables.

This script reads existing downstream summary CSVs only. It does not recompute
Cliff's delta, support status, bootstrap layers, or Stage 2 estimates.
"""

from __future__ import annotations

import csv
import math
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "MANUSCRIPT_RESULTS_BUNDLE_20260622_SUPPLEMENTARY_RESULTS_DOWNSTREAM"
TABLE_DIR = OUT_DIR / "tables"

Q2_PAIRWISE = ROOT / (
    "out_sensitivity/"
    "Sensitivity_RobustnessMatrix_Q2Cliff_AllOnly_NoAnnual_20260609_165018_allonly_noannual/"
    "summaries/robustness_area/robustness_area_summary.csv"
)
Q1_RAW_PERIOD = ROOT / (
    "out_sensitivity/"
    "Sensitivity_RawAllQ1Periods_CyanMagenta_Retention_20260615_122553/"
    "summaries/robustness_area/robustness_area_summary.csv"
)
GOOGLE_DOC_OVERLAY_SUMMARY = ROOT / (
    "out_sensitivity/"
    "Sensitivity_RobustnessMatrix_Q2Cliff_AllOnly_NoAnnual_20260609_165018_allonly_noannual/"
    "summaries/robustness_area/overlay_drafts/formula_brite_combo6yr_overlay_6panel_summary.csv"
)
MEDIUM_LARGE_OVERLAY = ROOT / (
    "out_sensitivity/"
    "MediumLargeEffectMasks_Q1Q2_4domain_20260615_20260615_2112_pmsm_annual_medium_large_countstyle_medium_large/"
    "summaries/medium_large_effect_masks_20260615/medium_large_effect_overlay_summary.csv"
)
LARGE_EFFECT_EXPANDED = ROOT / (
    "out_sensitivity/"
    "LargeEffectOverlay_Q1Q2_AllPMSM_4domain_RawMainOnly_20260615_101835_rawonly_large_effect_overlay/"
    "summaries/robustness_area/overlay_expanded_q1q2_all_pmsm_20260615/large_effect_overlay_expanded_summary.csv"
)

FOUR_DOMAIN_SET = {"Formula", "Brite", "Disease_NE", "Network"}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, fieldnames: list[str], rows: Iterable[dict[str, object]]) -> None:
    with path.open("w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def ffloat(value: str | int | float | None) -> float | None:
    if value is None or value == "":
        return None
    try:
        x = float(value)
    except ValueError:
        return None
    if math.isnan(x):
        return None
    return x


def fint(value: str | int | float | None) -> int | str:
    x = ffloat(value)
    if x is None:
        return ""
    return int(round(x))


def f3(value: str | int | float | None) -> str:
    x = ffloat(value)
    if x is None:
        return ""
    return f"{x:.3f}"


def pct(numerator: str | int | float | None, denominator: str | int | float | None) -> str:
    n = ffloat(numerator)
    d = ffloat(denominator)
    if n is None or d in (None, 0):
        return ""
    return f"{n / d:.3f}"


def comparison_family_label(row: dict[str, str], source_family: str) -> str:
    if source_family == "q1_raw_period":
        return "measurement_period_raw_q1"
    fam = row.get("Comparison_Family", "")
    if fam == "target_sample_coverage":
        return "target_sample_coverage_q2"
    if fam == "monte_carlo_seed":
        return "monte_carlo_seed_q2"
    if fam == "measurement_period":
        return "measurement_period_q2"
    if fam == "intensity_shape_within_stratum":
        return "intensity_shape_q2"
    if fam == "baseline_preservation":
        return "baseline_preservation_q2"
    return fam or source_family


def build_pairwise_rows() -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for source_path, source_family in [(Q2_PAIRWISE, "q2_all"), (Q1_RAW_PERIOD, "q1_raw_period")]:
        for row in read_csv(source_path):
            domain = row["Domain"]
            if domain not in FOUR_DOMAIN_SET:
                continue
            threshold = row["Threshold"]
            if threshold not in {"0.33", "0.474"}:
                continue
            out = {
                "Source_Family": source_family,
                "Overlay_Quantification_Family": comparison_family_label(row, source_family),
                "Comparison_Label": row["Comparison_Label"],
                "Domain": domain,
                "Subset": row["Subset"],
                "q": row["Q_Label"],
                "Threshold": threshold,
                "Direction": row["Direction"],
                "Paired_Cells": fint(row["Paired_Cells"]),
                "Baseline_Large_Cells": fint(row["Baseline_Large_Cells"]),
                "Variant_Large_Cells": fint(row["Variant_Large_Cells"]),
                "Intersection_Cells": fint(row["Intersection_Cells"]),
                "Union_Cells": fint(row["Union_Cells"]),
                "Retention": f3(row["Retention"]),
                "Jaccard": f3(row["Jaccard"]),
                "MedianAbsShift_InBaselineMask": f3(row["MedianAbsShift_InBaselineMask"]),
                "DirectionalAgreement_AllPaired": f3(row["DirectionalAgreement_AllPaired"]),
                "Baseline_Run": row.get("Baseline_Run", ""),
                "Variant_Run": row.get("Variant_Run", ""),
            }
            rows.append(out)
    rows.sort(
        key=lambda r: (
            str(r["Overlay_Quantification_Family"]),
            str(r["Domain"]),
            str(r["Subset"]),
            str(r["q"]),
            str(r["Threshold"]),
            str(r["Direction"]),
            str(r["Comparison_Label"]),
        )
    )
    return rows


def build_consensus_rows() -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []

    for row in read_csv(GOOGLE_DOC_OVERLAY_SUMMARY):
        if row["Domain"] not in {"Formula", "Brite"}:
            continue
        rows.append(
            {
                "Source_Family": "google_doc_figure5_large_only_06_2",
                "Overlay_ID": row["Overlay_Group"],
                "Overlay_Title": row["Overlay_Title"],
                "Domain": row["Domain"],
                "Subset": "All",
                "q": "q2",
                "Threshold_Definition": "large_only_0.474",
                "Direction": row["Direction"],
                "Scenario_Count": row["Scenario_Count"],
                "Grid_Cells": row["Grid_Cells"],
                "Any_Effect_Cells": row["Cells_Large_In_At_Least_One_Scenario"],
                "AllScenario_StatusOk_Cells": row["All_Scenario_Consensus_Cells_StatusOk"],
                "AllScenario_RegardlessStatus_Cells": row["Cells_Large_In_All_Scenarios_Regardless_Status"],
                "Missing_NotOk_Cells": row["Cells_Any_Status_Not_Ok_or_Missing"],
                "Consensus_Fraction_of_Any": pct(
                    row["All_Scenario_Consensus_Cells_StatusOk"],
                    row["Cells_Large_In_At_Least_One_Scenario"],
                ),
                "Any_Fraction_of_Grid": pct(row["Cells_Large_In_At_Least_One_Scenario"], row["Grid_Cells"]),
                "Scenario_List": "raw_0.965,TSC_0.935,TSC_0.990,seed_alt_98765"
                if row["Overlay_Group"] == "stability_tsc_seed"
                else "",
            }
        )

    for row in read_csv(MEDIUM_LARGE_OVERLAY):
        if row["Domain"] not in FOUR_DOMAIN_SET:
            continue
        rows.append(
            {
                "Source_Family": "supplementary_medium_large_overlay_20260622",
                "Overlay_ID": row["Overlay_ID"],
                "Overlay_Title": row["Overlay_ID"],
                "Domain": row["Domain"],
                "Subset": row["Subset"],
                "q": row["Q_Label"],
                "Threshold_Definition": "medium_or_larger_0.33_and_large_0.474",
                "Direction": row["Direction"],
                "Scenario_Count": row["Scenario_Count"],
                "Grid_Cells": row["Grid_Cells"],
                "Any_Effect_Cells": row["Any_MediumOrLarger_Cells"],
                "AllScenario_StatusOk_Cells": row["AllScenario_MediumOrLarger_StatusOk_Cells"],
                "AllScenario_RegardlessStatus_Cells": "",
                "Missing_NotOk_Cells": row["Any_Missing_NotOk_Cells"],
                "Consensus_Fraction_of_Any": pct(
                    row["AllScenario_MediumOrLarger_StatusOk_Cells"],
                    row["Any_MediumOrLarger_Cells"],
                ),
                "Any_Fraction_of_Grid": pct(row["Any_MediumOrLarger_Cells"], row["Grid_Cells"]),
                "Scenario_List": row["Scenario_List"],
            }
        )

    for row in read_csv(LARGE_EFFECT_EXPANDED):
        if row["Domain"] not in FOUR_DOMAIN_SET:
            continue
        rows.append(
            {
                "Source_Family": "supplementary_large_effect_measurement_period_20260615",
                "Overlay_ID": row["Overlay_Family"],
                "Overlay_Title": row["Overlay_Family"],
                "Domain": row["Domain"],
                "Subset": row["Subset"],
                "q": row["Q_Label"],
                "Threshold_Definition": "large_only_0.474",
                "Direction": row["Direction"],
                "Scenario_Count": row["Scenario_Count"],
                "Grid_Cells": row["Grid_Cells"],
                "Any_Effect_Cells": row["Any_Large_Cells"],
                "AllScenario_StatusOk_Cells": row["Consensus_Large_StatusOk_Cells"],
                "AllScenario_RegardlessStatus_Cells": "",
                "Missing_NotOk_Cells": row["Any_Missing_NotOk_Cells"],
                "Consensus_Fraction_of_Any": pct(row["Consensus_Large_StatusOk_Cells"], row["Any_Large_Cells"]),
                "Any_Fraction_of_Grid": pct(row["Any_Large_Cells"], row["Grid_Cells"]),
                "Scenario_List": row["Scenario_List"],
            }
        )

    rows.sort(
        key=lambda r: (
            str(r["Source_Family"]),
            str(r["Overlay_ID"]),
            str(r["Domain"]),
            str(r["Subset"]),
            str(r["q"]),
            str(r["Direction"]),
        )
    )
    return rows


def markdown_table(headers: list[str], rows: list[dict[str, object]], limit: int | None = None) -> str:
    use_rows = rows if limit is None else rows[:limit]
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in use_rows:
        vals = [str(row.get(h, "")).replace("|", "/") for h in headers]
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def write_markdown(pairwise_rows: list[dict[str, object]], consensus_rows: list[dict[str, object]]) -> None:
    pair_headers = [
        "Overlay_Quantification_Family",
        "Comparison_Label",
        "Domain",
        "Subset",
        "q",
        "Threshold",
        "Direction",
        "Baseline_Large_Cells",
        "Variant_Large_Cells",
        "Intersection_Cells",
        "Union_Cells",
        "Retention",
        "Jaccard",
        "MedianAbsShift_InBaselineMask",
        "DirectionalAgreement_AllPaired",
    ]
    consensus_headers = [
        "Source_Family",
        "Overlay_ID",
        "Domain",
        "Subset",
        "q",
        "Threshold_Definition",
        "Direction",
        "Scenario_Count",
        "Any_Effect_Cells",
        "AllScenario_StatusOk_Cells",
        "Missing_NotOk_Cells",
        "Consensus_Fraction_of_Any",
        "Any_Fraction_of_Grid",
    ]
    text = []
    text.append("# Full Retention/Jaccard and overlay-consensus quantification tables")
    text.append("")
    text.append("These tables summarize existing downstream overlay quantification outputs. They do not rerun the estimator or create independent biological replicates.")
    text.append("")
    text.append("## Pairwise Retention/Jaccard table")
    text.append("")
    text.append("Retention is the intersection of baseline and variant same-sign effect masks divided by the baseline effect-mask cells. Jaccard is the same-sign intersection divided by the union. Positive and negative masks are kept separate.")
    text.append("")
    text.append(markdown_table(pair_headers, pairwise_rows))
    text.append("")
    text.append("## Multi-scenario overlay consensus table")
    text.append("")
    text.append("These rows quantify multi-scenario overlays using all-scenario consensus cells and any-effect cells. They are not labeled Retention because no single pairwise baseline is defined for these multi-scenario overlays.")
    text.append("")
    text.append(markdown_table(consensus_headers, consensus_rows))
    text.append("")
    (TABLE_DIR / "retention_jaccard_full_pairwise_20260622.md").write_text("\n".join(text), encoding="utf-8")

    consensus_only = [
        "# Multi-scenario overlay consensus quantification",
        "",
        "This table is the multi-scenario companion to the Retention/Jaccard pairwise table.",
        "",
        markdown_table(consensus_headers, consensus_rows),
        "",
    ]
    (TABLE_DIR / "overlay_consensus_quantification_full_20260622.md").write_text(
        "\n".join(consensus_only), encoding="utf-8"
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def add_table(doc: Document, headers: list[str], rows: list[dict[str, object]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_layout_fixed(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
    for row in rows:
        tr = table.add_row()
        set_cant_split(tr)
        for i, header in enumerate(headers):
            cell = tr.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = str(row.get(header, ""))
            for p in cell.paragraphs:
                if header in {"Comparison_Label", "Source_Family", "Overlay_ID", "Threshold_Definition"}:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(7)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10 if style is None else 11)


def build_docx(pairwise_rows: list[dict[str, object]], consensus_rows: list[dict[str, object]]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Supplementary Table: Retention/Jaccard and overlay-consensus quantification")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)

    add_paragraph(
        doc,
        "This table bundle quantifies existing downstream overlay figures. Retention and Jaccard are pairwise mask-overlap metrics; multi-scenario overlays are reported using consensus and any-effect counts instead of being relabeled as pairwise Retention.",
    )

    doc.add_heading("Table S1. Pairwise Retention/Jaccard for four manuscript domains", level=1)
    add_paragraph(
        doc,
        "Rows are filtered to Formula, BRITE, Disease_NE, and Network. Positive and negative masks are separated. The full CSV includes run-path provenance columns.",
    )
    pair_headers = [
        "Overlay_Quantification_Family",
        "Comparison_Label",
        "Domain",
        "q",
        "Threshold",
        "Direction",
        "Baseline_Large_Cells",
        "Variant_Large_Cells",
        "Intersection_Cells",
        "Union_Cells",
        "Retention",
        "Jaccard",
        "DirectionalAgreement_AllPaired",
    ]
    add_table(doc, pair_headers, pairwise_rows, [1.25, 2.35, 0.8, 0.35, 0.55, 0.55, 0.65, 0.65, 0.65, 0.65, 0.55, 0.55, 0.75])

    doc.add_page_break()
    doc.add_heading("Table S2. Multi-scenario overlay consensus quantification", level=1)
    add_paragraph(
        doc,
        "These rows cover the Google Docs Figure 5 large-only overlay summary and the overlay families included in the 2026-06-22 Supplementary Results bundle. Consensus_Fraction_of_Any is all-scenario status-ok consensus cells divided by cells marked in at least one scenario.",
    )
    consensus_headers = [
        "Source_Family",
        "Overlay_ID",
        "Domain",
        "Subset",
        "q",
        "Threshold_Definition",
        "Direction",
        "Scenario_Count",
        "Any_Effect_Cells",
        "AllScenario_StatusOk_Cells",
        "Missing_NotOk_Cells",
        "Consensus_Fraction_of_Any",
    ]
    add_table(
        doc,
        consensus_headers,
        consensus_rows,
        [1.55, 1.25, 0.75, 0.45, 0.35, 1.25, 0.55, 0.55, 0.65, 0.75, 0.65, 0.75],
    )

    doc_path = OUT_DIR / "Supplementary_Table_Retention_Jaccard_overlay_quantification_20260622.docx"
    doc.save(doc_path)
    return doc_path


def main() -> None:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    pairwise_rows = build_pairwise_rows()
    consensus_rows = build_consensus_rows()

    pair_fields = [
        "Source_Family",
        "Overlay_Quantification_Family",
        "Comparison_Label",
        "Domain",
        "Subset",
        "q",
        "Threshold",
        "Direction",
        "Paired_Cells",
        "Baseline_Large_Cells",
        "Variant_Large_Cells",
        "Intersection_Cells",
        "Union_Cells",
        "Retention",
        "Jaccard",
        "MedianAbsShift_InBaselineMask",
        "DirectionalAgreement_AllPaired",
        "Baseline_Run",
        "Variant_Run",
    ]
    consensus_fields = [
        "Source_Family",
        "Overlay_ID",
        "Overlay_Title",
        "Domain",
        "Subset",
        "q",
        "Threshold_Definition",
        "Direction",
        "Scenario_Count",
        "Grid_Cells",
        "Any_Effect_Cells",
        "AllScenario_StatusOk_Cells",
        "AllScenario_RegardlessStatus_Cells",
        "Missing_NotOk_Cells",
        "Consensus_Fraction_of_Any",
        "Any_Fraction_of_Grid",
        "Scenario_List",
    ]

    write_csv(TABLE_DIR / "retention_jaccard_full_pairwise_20260622.csv", pair_fields, pairwise_rows)
    write_csv(TABLE_DIR / "overlay_consensus_quantification_full_20260622.csv", consensus_fields, consensus_rows)
    write_markdown(pairwise_rows, consensus_rows)
    doc_path = build_docx(pairwise_rows, consensus_rows)

    print(f"pairwise_rows={len(pairwise_rows)}")
    print(f"consensus_rows={len(consensus_rows)}")
    print(f"docx={doc_path}")


if __name__ == "__main__":
    main()
